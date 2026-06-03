from __future__ import annotations

import io
from dataclasses import dataclass
from pathlib import Path

from docx import Document as DocxDocument
from docx.document import Document as DocumentType
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.document_assets import format_image_block
from app.loaders.errors import LoaderError

MIN_IMAGE_BYTES = 500

@dataclass
class DocxImageRef:
    index: int
    image_id: str
    rel_id: str
    data: bytes
    mime_type: str


def load_docx(path: Path) -> str:
    try:
        document = DocxDocument(str(path))
        text = compose_docx_text(document, ocr_text_by_id={}).strip()
    except Exception as exc:
        raise LoaderError("extraction_failed") from exc

    if not text:
        raise LoaderError("extraction_failed")
    return text


def build_docx_image_id_map(document: DocumentType) -> dict[str, str]:
    mapping: dict[str, str] = {}
    index = 1
    for rel_id in _iter_docx_image_rel_ids(document):
        if rel_id in mapping:
            continue
        blob, _ = _blob_for_rel_id(document, rel_id)
        if blob is None or len(blob) < MIN_IMAGE_BYTES:
            continue
        mapping[rel_id] = f"img_{index:03d}"
        index += 1
    return mapping


def extract_docx_images_ordered(content: bytes) -> list[DocxImageRef]:
    document = DocxDocument(io.BytesIO(content))
    rel_map = build_docx_image_id_map(document)
    refs: list[DocxImageRef] = []
    for rel_id, image_id in rel_map.items():
        blob, mime_type = _blob_for_rel_id(document, rel_id)
        if blob is None:
            continue
        index = int(image_id.split("_")[1])
        refs.append(
            DocxImageRef(
                index=index,
                image_id=image_id,
                rel_id=rel_id,
                data=blob,
                mime_type=mime_type,
            )
        )
    refs.sort(key=lambda item: item.index)
    return refs


def compose_docx_text(
    document: DocumentType,
    *,
    ocr_text_by_id: dict[str, str],
    include_unprocessed_placeholders: bool = False,
) -> str:
    rel_map = build_docx_image_id_map(document)
    parts: list[str] = []

    for block in _iter_block_items(document):
        if isinstance(block, Paragraph):
            paragraph_text = _compose_paragraph(
                block,
                rel_map,
                ocr_text_by_id,
                include_unprocessed_placeholders,
            )
            if paragraph_text:
                parts.append(paragraph_text)
        elif isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph_text = _compose_paragraph(
                            paragraph,
                            rel_map,
                            ocr_text_by_id,
                            include_unprocessed_placeholders,
                        )
                        if paragraph_text:
                            parts.append(paragraph_text)

    return "\n\n".join(part for part in parts if part.strip())


def _compose_paragraph(
    paragraph: Paragraph,
    rel_map: dict[str, str],
    ocr_text_by_id: dict[str, str],
    include_unprocessed_placeholders: bool,
) -> str:
    chunks: list[str] = []

    for run in paragraph.runs:
        rel_ids = _rel_ids_in_run(run._element)
        if rel_ids:
            if run.text.strip():
                chunks.append(run.text.strip())
            for rel_id in rel_ids:
                image_id = rel_map.get(rel_id)
                if not image_id:
                    continue
                ocr_text = ocr_text_by_id.get(image_id)
                if ocr_text:
                    chunks.append(format_image_block(image_id=image_id, page=None, transcription=ocr_text))
                elif include_unprocessed_placeholders:
                    chunks.append(f'[BILD id="{image_id}" status="nicht_verarbeitet"]')
        elif run.text:
            chunks.append(run.text)

    return " ".join(part.strip() for part in chunks if part.strip()).strip()


def _iter_docx_image_rel_ids(document: DocumentType) -> list[str]:
    rel_ids: list[str] = []
    seen: set[str] = set()
    for block in _iter_block_items(document):
        paragraphs = [block] if isinstance(block, Paragraph) else [p for row in block.rows for cell in row.cells for p in cell.paragraphs]
        for paragraph in paragraphs:
            for rel_id in _rel_ids_in_paragraph(paragraph):
                if rel_id not in seen:
                    seen.add(rel_id)
                    rel_ids.append(rel_id)
    return rel_ids


def _iter_block_items(document: DocumentType):
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _rel_ids_in_paragraph(paragraph: Paragraph) -> list[str]:
    rel_ids: list[str] = []
    for run in paragraph.runs:
        rel_ids.extend(_rel_ids_in_run(run._element))
    return rel_ids


def _rel_ids_in_run(element) -> list[str]:
    rel_ids: list[str] = []
    for blip in element.findall(".//" + qn("a:blip")):
        embed = blip.get(qn("r:embed"))
        if embed:
            rel_ids.append(embed)
    return rel_ids


def _blob_for_rel_id(document: DocumentType, rel_id: str) -> tuple[bytes | None, str]:
    rel = document.part.rels.get(rel_id)
    if rel is None or "image" not in rel.target_ref:
        return None, "image/png"
    blob = rel.target_part.blob
    return blob, _guess_mime(blob, rel.target_ref)


def _guess_mime(data: bytes, name: str = "") -> str:
    lowered = name.lower()
    if lowered.endswith(".png") or data.startswith(b"\x89PNG\r\n\x1a\n"):
        return "image/png"
    if lowered.endswith(".jpg") or lowered.endswith(".jpeg") or data.startswith(b"\xff\xd8"):
        return "image/jpeg"
    if lowered.endswith(".gif") or data.startswith(b"GIF87a") or data.startswith(b"GIF89a"):
        return "image/gif"
    if lowered.endswith(".webp") or data[0:4] == b"RIFF" and data[8:12] == b"WEBP":
        return "image/webp"
    return "image/png"
