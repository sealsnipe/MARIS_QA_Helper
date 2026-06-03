from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Inches
from PIL import Image

from app.document_assets import SavedDocumentImage, format_image_block
from app.loaders.docx_content import build_docx_image_id_map, compose_docx_text, extract_docx_images_ordered
from app.loaders.vision_ocr import VisionOcrResult, compose_docx_with_vision


def _png_bytes(width: int = 200, height: int = 200) -> bytes:
    buffer = BytesIO()
    Image.new("RGB", (width, height), color="red").save(buffer, format="PNG")
    return buffer.getvalue()


def _make_docx_with_two_images(tmp_path: Path) -> Path:
    path = tmp_path / "mixed.docx"
    png1 = tmp_path / "a.png"
    png2 = tmp_path / "b.png"
    png1.write_bytes(_png_bytes())
    png2.write_bytes(_png_bytes(220, 220))

    document = DocxDocument()
    document.add_paragraph("Einleitung vor Bild eins.")
    document.add_paragraph("").add_run().add_picture(str(png1), width=Inches(1.0))
    document.add_paragraph("Text zwischen den Bildern.")
    document.add_paragraph("").add_run().add_picture(str(png2), width=Inches(1.0))
    document.add_paragraph("Abschluss nach Bild zwei.")
    document.save(path)
    return path


def test_extract_docx_images_ordered_follows_document_flow(tmp_path: Path) -> None:
    path = _make_docx_with_two_images(tmp_path)
    refs = extract_docx_images_ordered(path.read_bytes())
    assert len(refs) == 2
    assert refs[0].image_id == "img_001"
    assert refs[1].image_id == "img_002"


def test_compose_docx_text_interleaves_ocr_blocks(tmp_path: Path) -> None:
    path = _make_docx_with_two_images(tmp_path)
    document = DocxDocument(str(path))
    rel_map = build_docx_image_id_map(document)
    assert len(rel_map) == 2

    ocr_by_id = {
        "img_001": "Inhalt von Bild eins",
        "img_002": "Inhalt von Bild zwei",
    }
    text = compose_docx_text(document, ocr_text_by_id=ocr_by_id)
    intro = text.find("Einleitung vor Bild eins.")
    block1 = text.find('[BILD id="img_001"]')
    middle = text.find("Text zwischen den Bildern.")
    block2 = text.find('[BILD id="img_002"]')
    outro = text.find("Abschluss nach Bild zwei.")

    assert intro != -1 and block1 != -1 and middle != -1 and block2 != -1 and outro != -1
    assert intro < block1 < middle < block2 < outro
    assert "Inhalt von Bild eins" in text
    assert format_image_block(image_id="img_002", page=None, transcription="x")[:20] in text or "Inhalt von Bild zwei" in text


def test_compose_docx_unprocessed_placeholder(tmp_path: Path) -> None:
    path = _make_docx_with_two_images(tmp_path)
    document = DocxDocument(str(path))
    text = compose_docx_text(document, ocr_text_by_id={}, include_unprocessed_placeholders=True)
    assert '[BILD id="img_001" status="nicht_verarbeitet"]' in text
    assert "Einleitung vor Bild eins." in text


def test_compose_docx_with_vision_interleaves_blocks(tmp_path: Path) -> None:
    path = _make_docx_with_two_images(tmp_path)
    blocks = [
        format_image_block(image_id="img_001", page=None, transcription="Inhalt von Bild eins"),
        format_image_block(image_id="img_002", page=None, transcription="Inhalt von Bild zwei"),
    ]
    result = VisionOcrResult(
        blocks=blocks,
        images_processed=2,
        images_failed=0,
        saved_images=[
            SavedDocumentImage(id="img_001", filename="img_001.png", page=None, mime_type="image/png"),
            SavedDocumentImage(id="img_002", filename="img_002.png", page=None, mime_type="image/png"),
        ],
    )

    text = compose_docx_with_vision(path, result)
    intro = text.find("Einleitung vor Bild eins.")
    block1 = text.find('[BILD id="img_001"]')
    middle = text.find("Text zwischen den Bildern.")
    block2 = text.find('[BILD id="img_002"]')
    outro = text.find("Abschluss nach Bild zwei.")

    assert intro != -1 and block1 != -1 and middle != -1 and block2 != -1 and outro != -1
    assert intro < block1 < middle < block2 < outro
