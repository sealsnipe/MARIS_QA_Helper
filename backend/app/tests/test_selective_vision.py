from __future__ import annotations

import os
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Inches
from PIL import Image

from app.loaders.vision_ocr import run_vision_ocr, save_embedded_images
from app.upload import inspect_upload, parse_transcribe_image_ids
from app.tests.conftest import create_customer


def _png_bytes(width: int = 240, height: int = 240) -> bytes:
    """Blocky random-color image so the downsampled 32x32 gray in _is_meaningful_image still
    has sufficient variance (>180 strict) and unique colors. Pure per-pixel os.urandom noise
    gets smoothed too much by LANCZOS resize and would be (incorrectly for test) filtered.
    Block structure survives downsampling and represents "realistic embedded figure".
    """
    im = Image.new("RGB", (width, height))
    block = 8
    for y in range(0, height, block):
        for x in range(0, width, block):
            r = int.from_bytes(os.urandom(1), "big")
            g = int.from_bytes(os.urandom(1), "big")
            b = int.from_bytes(os.urandom(1), "big")
            for yy in range(block):
                for xx in range(block):
                    if y + yy < height and x + xx < width:
                        im.putpixel((x + xx, y + yy), (r, g, b))
    buffer = BytesIO()
    im.save(buffer, format="PNG")
    return buffer.getvalue()


def _make_docx_with_two_images(tmp_path: Path) -> bytes:
    png1 = tmp_path / "a.png"
    png2 = tmp_path / "b.png"
    png1.write_bytes(_png_bytes())
    png2.write_bytes(_png_bytes(220, 220))

    document = DocxDocument()
    document.add_paragraph("Einleitung vor Bild eins.")
    document.add_paragraph("").add_run().add_picture(str(png1), width=Inches(1.0))
    document.add_paragraph("Text zwischen den Bildern.")
    document.add_paragraph("").add_run().add_picture(str(png2), width=Inches(1.0))
    document.add_paragraph("Abschluss nach Bild zwei.")

    path = tmp_path / "mixed.docx"
    document.save(path)
    return path.read_bytes()


def test_inspect_upload_returns_image_previews(db_session, tmp_path: Path) -> None:
    create_customer(db_session, "vision-test", "Vision Test")
    content = _make_docx_with_two_images(tmp_path)
    payload = inspect_upload(db_session, "vision-test", content, "mixed.docx")
    assert payload["image_count"] == 2
    assert len(payload["images"]) == 2
    assert payload["images"][0]["id"] == "img_001"
    assert payload["images"][0]["preview_data_url"].startswith("data:image/")


def test_parse_transcribe_image_ids_accepts_json_and_csv() -> None:
    assert parse_transcribe_image_ids('["img_001","img_003"]') == {"img_001", "img_003"}
    assert parse_transcribe_image_ids("img_001, img_002") == {"img_001", "img_002"}
    assert parse_transcribe_image_ids("invalid, img_0042") == set()


def test_run_vision_ocr_only_transcribes_selected_ids(tmp_path: Path, monkeypatch) -> None:
    docx_path = tmp_path / "mixed.docx"
    docx_path.write_bytes(_make_docx_with_two_images(tmp_path))
    assets_dir = tmp_path / "images"
    saved = save_embedded_images(docx_path, ".docx", assets_dir)

    calls: list[str] = []

    def fake_transcribe(image_data, mime_type, *, prompt):
        calls.append("x")
        return "OCR Text"

    monkeypatch.setattr("app.loaders.vision_ocr.transcribe_image", fake_transcribe)

    result = run_vision_ocr(
        docx_path,
        ".docx",
        assets_dir=assets_dir,
        transcribe_ids={"img_002"},
        saved_images=saved,
    )
    assert result.images_processed == 1
    assert len(result.blocks) == 1
    assert 'id="img_002"' in result.blocks[0]
    assert len(calls) == 1
    assert result.saved_images[0].transcribed is False
    assert result.saved_images[1].transcribed is True
